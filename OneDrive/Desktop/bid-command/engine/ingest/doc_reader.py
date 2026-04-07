"""Document reader for PDF, DOCX, XLSX, and plain-text files.

Provides a single public function, ``read_document(path)``, that dispatches to
the appropriate reader based on file suffix and returns a ``DocContent`` object.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


# ---------------------------------------------------------------------------
# Public dataclass
# ---------------------------------------------------------------------------

@dataclass
class DocContent:
    """Structured result of reading a single document."""

    filename: str
    file_type: str      # "pdf", "docx", "xlsx", "txt", "csv", "html", "unknown"
    text: str
    page_count: int
    file_path: Path


# ---------------------------------------------------------------------------
# Private reader implementations
# ---------------------------------------------------------------------------

def _read_pdf(path: Path) -> DocContent:
    """Read a PDF with PyMuPDF (fitz), labelling each page."""
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    total = doc.page_count
    parts: list[str] = []
    for i, page in enumerate(doc):
        header = f"--- PAGE {i + 1} of {total} [{path.name}] ---"
        text = page.get_text()
        parts.append(f"{header}\n{text}")
    doc.close()

    return DocContent(
        filename=path.name,
        file_type="pdf",
        text="\n".join(parts),
        page_count=total,
        file_path=path,
    )


def _read_xlsx(path: Path) -> DocContent:
    """Read an XLSX with openpyxl, outputting each sheet as pipe-separated rows."""
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        header = f"--- SHEET: {ws.title} [{path.name}] ---"
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            # Convert each cell to string, replacing None with empty string
            rows.append(" | ".join("" if cell is None else str(cell) for cell in row))
        parts.append(header + "\n" + "\n".join(rows))

    text = "\n\n".join(parts)
    # XLSX doesn't have a page count concept; use sheet count
    page_count = len(wb.worksheets)

    return DocContent(
        filename=path.name,
        file_type="xlsx",
        text=text,
        page_count=page_count,
        file_path=path,
    )


def _read_docx(path: Path) -> DocContent:
    """Read a DOCX with python-docx, extracting paragraphs and table text."""
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    text = "\n".join(parts)
    # Estimate page count at ~3000 chars per page (minimum 1 if there's content)
    char_count = len(text)
    page_count = max(1, (char_count + 2999) // 3000) if text else 0

    return DocContent(
        filename=path.name,
        file_type="docx",
        text=text,
        page_count=page_count,
        file_path=path,
    )


def _read_text(path: Path, file_type: str) -> DocContent:
    """Read a plain-text file (txt, csv, html) as UTF-8."""
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except OSError:
        text = ""

    # Estimate pages at ~3000 chars each
    char_count = len(text)
    page_count = max(1, (char_count + 2999) // 3000) if text else 0

    return DocContent(
        filename=path.name,
        file_type=file_type,
        text=text,
        page_count=page_count,
        file_path=path,
    )


def _read_unknown(path: Path) -> DocContent:
    """Return an empty DocContent for unsupported file types."""
    return DocContent(
        filename=path.name,
        file_type="unknown",
        text="",
        page_count=0,
        file_path=path,
    )


# ---------------------------------------------------------------------------
# Dispatch table
# ---------------------------------------------------------------------------

def _dispatch_pdf(path: Path) -> DocContent:
    return _read_pdf(path)

def _dispatch_xlsx(path: Path) -> DocContent:
    return _read_xlsx(path)

def _dispatch_docx(path: Path) -> DocContent:
    return _read_docx(path)

def _dispatch_txt(path: Path) -> DocContent:
    return _read_text(path, "txt")

def _dispatch_csv(path: Path) -> DocContent:
    return _read_text(path, "csv")

def _dispatch_html(path: Path) -> DocContent:
    return _read_text(path, "html")


_SUFFIX_DISPATCH: dict[str, object] = {
    ".pdf":  _dispatch_pdf,
    ".xlsx": _dispatch_xlsx,
    ".xls":  _dispatch_xlsx,
    ".docx": _dispatch_docx,
    ".doc":  _dispatch_docx,
    ".txt":  _dispatch_txt,
    ".csv":  _dispatch_csv,
    ".html": _dispatch_html,
    ".htm":  _dispatch_html,
}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def read_document(path: Path) -> DocContent:
    """Read *path* and return a :class:`DocContent` with extracted text.

    The file type is determined by the file suffix (case-insensitive).
    Unsupported types return a ``DocContent`` with ``file_type="unknown"``,
    empty ``text``, and ``page_count=0``.

    Parameters
    ----------
    path:
        Absolute or relative path to the document file.

    Returns
    -------
    DocContent
        Parsed document content and metadata.
    """
    suffix = path.suffix.lower()
    reader = _SUFFIX_DISPATCH.get(suffix)
    if reader is None:
        return _read_unknown(path)
    return reader(path)  # type: ignore[operator]
